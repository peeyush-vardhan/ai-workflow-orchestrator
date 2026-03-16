"""Export blueprint — download workflow output as MD, PDF, or DOCX."""
import io
import re
from flask import Blueprint, Response, abort, current_app, request

export_bp = Blueprint("export", __name__)

_SLUG_RE = re.compile(r"[^a-z0-9]+")


def _slug(name: str) -> str:
    return _SLUG_RE.sub("-", name.lower()).strip("-") or "workflow"


def _build_full_markdown(state) -> str:
    """Assemble a single Markdown document from all completed task outputs."""
    lines = [f"# {state.dag.workflow_name}", "", f"**Input:** {state.user_input}", "", "---", ""]
    for task in state.dag.get_execution_order():
        if task.output:
            agent = task.agent_type.value if hasattr(task.agent_type, "value") else task.agent_type
            lines += [
                f"## {task.description}",
                f"*Agent: {agent} · Task: {task.id}*",
                "",
                task.output,
                "",
                "---",
                "",
            ]
    return "\n".join(lines)


def _strip_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


@export_bp.route("/api/workflows/<workflow_id>/export", methods=["GET"])
def export_workflow(workflow_id: str):
    store = current_app.config["STORE"]
    state = store.load(workflow_id)
    if not state or not state.dag:
        abort(404, description="Workflow not found")

    fmt = request.args.get("format", "md").lower()
    name = state.dag.workflow_name
    slug = _slug(name)
    content = _build_full_markdown(state)

    if fmt == "md":
        return Response(
            content,
            mimetype="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{slug}.md"'},
        )

    if fmt == "pdf":
        return _export_pdf(content, name, slug)

    if fmt == "docx":
        return _export_docx(state, name, slug)

    abort(400, description=f"Unknown format '{fmt}'. Use md, pdf, or docx.")


def _export_pdf(content: str, title: str, slug: str) -> Response:
    try:
        from fpdf import FPDF as _FPDF  # fpdf2 package installs as 'fpdf'
    except ImportError:
        abort(500, description="PDF export requires: pip install fpdf2")

    pdf = _FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin

    for line in content.split("\n"):
        clean = _strip_inline(line)

        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 18)
            pdf.set_text_color(0, 160, 220)
            pdf.multi_cell(effective_w, 9, clean[2:])
            pdf.ln(2)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(30, 30, 80)
            pdf.ln(4)
            pdf.multi_cell(effective_w, 8, clean[3:])
            pdf.ln(1)
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(60, 60, 100)
            pdf.ln(3)
            pdf.multi_cell(effective_w, 7, clean[4:])
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(effective_w, 6, f"  \u2022 {clean[2:]}")
        elif line.strip() == "---":
            pdf.ln(2)
            pdf.set_draw_color(180, 180, 200)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_w, pdf.get_y())
            pdf.ln(2)
        elif line.strip() == "":
            pdf.ln(3)
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(40, 40, 40)
            if clean.strip():
                pdf.multi_cell(effective_w, 6, clean)

    pdf_bytes = bytes(pdf.output())
    return Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{slug}.pdf"'},
    )


def _export_docx(state, name: str, slug: str) -> Response:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        abort(500, description="DOCX export requires: pip install python-docx")

    doc = Document()
    doc.core_properties.title = name

    doc.add_heading(name, 0)
    doc.add_paragraph(f"Input: {state.user_input}")

    for task in state.dag.get_execution_order():
        if not task.output:
            continue
        agent = task.agent_type.value if hasattr(task.agent_type, "value") else task.agent_type
        doc.add_heading(task.description, level=1)
        doc.add_paragraph(f"Agent: {agent}  |  Task: {task.id}").italic = True

        for line in task.output.split("\n"):
            clean = _strip_inline(line)
            if line.startswith("## "):
                doc.add_heading(clean[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(clean[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(clean[2:], style="List Bullet")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                if clean.strip():
                    doc.add_paragraph(clean)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{slug}.docx"'},
    )
